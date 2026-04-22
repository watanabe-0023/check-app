"""収益管理表 月次チェックアプリ。"""
import io
from pathlib import Path

import pandas as pd
import streamlit as st

from core import loaders, reconcilers
from core.schema import RuleResult

st.set_page_config(page_title="収益管理表 月次チェック", layout="wide", page_icon="📊")

# パスワード認証
PASSWORD = "978init"
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False

if not st.session_state["authenticated"]:
    st.title("🔒 収益管理表 月次チェックアプリ")
    pw = st.text_input("パスワードを入力してください", type="password")
    if st.button("ログイン"):
        if pw == PASSWORD:
            st.session_state["authenticated"] = True
            st.rerun()
        else:
            st.error("パスワードが違います")
    st.stop()

st.title("📊 収益管理表 月次チェックアプリ")
st.caption("init株式会社 | 会計事務所向け 自動突合ツール")

tab1, tab2, tab3 = st.tabs(["① アップロード", "② 突合結果", "③ レポート出力"])

# ---------------------------------------------------------------------------
# タブ1: アップロード
# ---------------------------------------------------------------------------
with tab1:
    st.subheader("ファイルアップロード")
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("**必須ファイル**")
        f_mgmt  = st.file_uploader("📁 収益管理表 (xlsx)", type=["xlsx"], key="mgmt")
        f_sales = st.file_uploader("📁 売上管理台帳 (xlsx)", type=["xlsx"], key="sales")
        f_exp   = st.file_uploader("📁 支出管理表 (xlsx)", type=["xlsx"], key="exp")
        f_bs    = st.file_uploader("📁 freee 試算表 BS (xlsx/pdf)", type=["xlsx", "pdf"], key="bs")
        f_pl    = st.file_uploader("📁 freee 月次推移表 PL (xlsx/pdf)", type=["xlsx", "pdf"], key="pl")

    with col2:
        st.markdown("**任意ファイル**")
        f_keihi = st.file_uploader("📁 経費台帳 (xlsx)", type=["xlsx"], key="keihi")
        f_biz   = st.file_uploader("📁 事業別支出管理 (xlsx)", type=["xlsx"], key="biz")

    st.divider()
    st.subheader("対象月選択")

    months = loaders.list_months(f_mgmt)
    if months:
        month = st.selectbox("対象月", options=months, index=len(months) - 1)
    else:
        month_input = st.text_input("対象月 (YYYY-MM形式)", placeholder="例: 2026-01")
        month = month_input.strip() if month_input else None
        if not months:
            st.info("収益管理表をアップロードすると対象月が自動検出されます。")

    st.divider()
    required_ok = all([f_mgmt, f_sales, f_exp, f_bs, f_pl]) and bool(month)
    if not required_ok:
        st.warning("必須ファイル（収益管理表・売上台帳・支出管理表・freee BS・freee PL）をすべてアップロードし、対象月を選択してください。")

    if st.button("▶️ チェック実行", disabled=not required_ok, type="primary", use_container_width=True):
        files = {
            "mgmt": f_mgmt, "sales": f_sales, "exp": f_exp,
            "keihi": f_keihi, "biz": f_biz, "bs": f_bs, "pl": f_pl,
        }
        with st.spinner("突合処理中..."):
            try:
                st.session_state["results"] = reconcilers.run_all(files, month)
                st.session_state["month"] = month
                st.success("✅ チェック完了！「② 突合結果」タブで確認してください。")
            except Exception as e:
                st.error(f"エラーが発生しました: {e}")
                import traceback
                st.code(traceback.format_exc())


# ---------------------------------------------------------------------------
# タブ2: 突合結果
# ---------------------------------------------------------------------------
with tab2:
    if "results" not in st.session_state:
        st.info("「① アップロード」タブでファイルをアップロードし、チェックを実行してください。")
        st.stop()

    data = st.session_state["results"]
    results: list[RuleResult] = data["results"]
    month_label = data["month"]
    irregulars = data.get("irregulars", [])
    hash_changes = data.get("hash_changes", {})

    st.subheader(f"突合結果 — {month_label}")

    # サマリカード
    counts = {"pass": 0, "warn": 0, "fail": 0, "info": 0}
    for r in results:
        counts[r.status] = counts.get(r.status, 0) + 1

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("総チェック数", len(results))
    c2.metric("✅ PASS", counts["pass"])
    c3.metric("⚠️ 要確認", counts["warn"])
    c4.metric("❌ NG", counts["fail"])
    c5.metric("ℹ️ 参考", counts["info"])

    st.divider()

    # フィルタ
    filter_status = st.multiselect(
        "ステータスでフィルタ",
        options=["pass", "warn", "fail", "info"],
        default=["warn", "fail", "info"],
        format_func=lambda x: {"pass": "✅ PASS", "warn": "⚠️ 要確認", "fail": "❌ NG", "info": "ℹ️ 参考"}[x],
    )

    filtered = [r for r in results if r.status in filter_status]

    # 差異一覧テーブル
    def status_icon(s):
        return {"pass": "✅", "warn": "⚠️", "fail": "❌", "info": "ℹ️"}.get(s, s)

    rows = []
    for r in filtered:
        rows.append({
            "ST": status_icon(r.status),
            "ID": r.rule_id,
            "チェック項目": r.label,
            "管理表値 (A)": f"{r.a_value:,.0f}" if r.a_value is not None else "—",
            "実績値 (B)": f"{r.b_value:,.0f}" if r.b_value is not None else "—",
            "差異 (A-B)": f"{r.diff:+,.0f}" if r.diff is not None else "—",
            "差異率": f"{r.diff_rate:.2%}" if r.diff_rate is not None else "—",
            "備考": r.note,
        })

    if rows:
        df_display = pd.DataFrame(rows)

        def color_row(row):
            status_map = {"✅": "background-color: #d4edda", "⚠️": "background-color: #fff3cd",
                          "❌": "background-color: #f8d7da", "ℹ️": "background-color: #d1ecf1"}
            color = status_map.get(row["ST"], "")
            return [color] * len(row)

        styled = df_display.style.apply(color_row, axis=1)
        st.dataframe(styled, use_container_width=True, hide_index=True)

        # 詳細パネル
        st.divider()
        st.subheader("詳細確認")
        selected_id = st.selectbox("確認したいルールIDを選択",
                                   options=[r.rule_id for r in filtered],
                                   format_func=lambda x: f"{x} — {next((r.label for r in filtered if r.rule_id == x), x)}")
        if selected_id:
            r = next(r for r in filtered if r.rule_id == selected_id)
            with st.expander(f"📋 {r.rule_id}: {r.label}", expanded=True):
                dc1, dc2, dc3 = st.columns(3)
                dc1.metric("管理表値 (A)", f"{r.a_value:,.0f}円" if r.a_value is not None else "—")
                dc2.metric("実績値 (B)", f"{r.b_value:,.0f}円" if r.b_value is not None else "—")
                dc3.metric("差異 (A-B)", f"{r.diff:+,.0f}円" if r.diff is not None else "—")
                st.markdown(f"**根拠A:** `{r.a_source}`")
                st.markdown(f"**根拠B:** `{r.b_source}`")
                if r.note:
                    st.info(r.note)
    else:
        st.success("フィルタ条件に一致する項目はありません。")

    # イレギュラー検知結果
    if irregulars or hash_changes:
        st.divider()
        st.subheader("🔍 台帳イレギュラー検知")

        if hash_changes:
            st.warning(f"⚠️ 前回実行から変更が検出されました: {list(hash_changes.keys())}")
            for k, v in hash_changes.items():
                st.caption(f"  {k}: {v['prev']} → {v['current']}")

        if irregulars:
            st.warning(f"⚠️ {len(irregulars)} 件のイレギュラーが検出されました。")
            irr_df = pd.DataFrame(irregulars)
            st.dataframe(irr_df[["type", "file", "cell", "description"]], use_container_width=True, hide_index=True)
        else:
            st.success("台帳イレギュラーは検出されませんでした。")


# ---------------------------------------------------------------------------
# タブ3: レポート出力
# ---------------------------------------------------------------------------
with tab3:
    if "results" not in st.session_state:
        st.info("先にチェックを実行してください。")
        st.stop()

    data = st.session_state["results"]
    results: list[RuleResult] = data["results"]
    month_label = data["month"]

    st.subheader("レポート出力")

    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("### Excel差異報告書")
        if st.button("📥 Excel出力", use_container_width=True):
            buf = _build_excel_report(results, month_label)
            st.download_button(
                "ダウンロード (xlsx)",
                data=buf,
                file_name=f"差異報告_{month_label}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

    with col_b:
        st.markdown("### Word差異サマリ")
        if st.button("📥 Word出力", use_container_width=True):
            buf = _build_word_report(results, month_label, data.get("irregulars", []))
            st.download_button(
                "ダウンロード (docx)",
                data=buf,
                file_name=f"差異サマリ_{month_label}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )


# ---------------------------------------------------------------------------
# レポート生成関数（タブ外で定義）
# ---------------------------------------------------------------------------

def _build_excel_report(results: list[RuleResult], month: str) -> bytes:
    import openpyxl
    from openpyxl.styles import PatternFill, Font, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "差異報告"

    # ヘッダ
    headers = ["ステータス", "ルールID", "チェック項目", "管理表値(A)", "実績値(B)", "差異(A-B)", "差異率", "根拠A", "根拠B", "備考"]
    fill_header = PatternFill("solid", fgColor="2F4F7F")
    font_header = Font(color="FFFFFF", bold=True)
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = fill_header
        cell.font = font_header
        cell.alignment = Alignment(horizontal="center")

    status_fills = {
        "pass": PatternFill("solid", fgColor="D4EDDA"),
        "warn": PatternFill("solid", fgColor="FFF3CD"),
        "fail": PatternFill("solid", fgColor="F8D7DA"),
        "info": PatternFill("solid", fgColor="D1ECF1"),
    }
    status_labels = {"pass": "PASS", "warn": "要確認", "fail": "NG", "info": "参考"}

    for row_idx, r in enumerate(results, 2):
        fill = status_fills.get(r.status, PatternFill())
        vals = [
            status_labels.get(r.status, r.status),
            r.rule_id,
            r.label,
            r.a_value,
            r.b_value,
            r.diff,
            f"{r.diff_rate:.2%}" if r.diff_rate is not None else "",
            r.a_source,
            r.b_source,
            r.note,
        ]
        for col_idx, val in enumerate(vals, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.fill = fill

    # 列幅調整
    widths = [10, 8, 30, 15, 15, 12, 10, 30, 30, 40]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # タイトル行追加
    ws.insert_rows(1)
    title_cell = ws.cell(row=1, column=1, value=f"収益管理表 差異報告書 — {month}")
    title_cell.font = Font(size=14, bold=True)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _build_word_report(results: list[RuleResult], month: str, irregulars: list) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # python-docx が未インストールの場合はシンプルなテキストファイルで代替
        lines = [f"収益管理表 差異サマリ — {month}\n"]
        counts = {"pass": 0, "warn": 0, "fail": 0, "info": 0}
        for r in results:
            counts[r.status] = counts.get(r.status, 0) + 1
        lines.append(f"合計: {len(results)}件 / PASS: {counts['pass']} / 要確認: {counts['warn']} / NG: {counts['fail']}\n\n")
        for r in results:
            lines.append(f"[{r.rule_id}] {r.label}: {r.status.upper()}\n")
            if r.diff is not None:
                lines.append(f"  差異: {r.diff:+,.0f}円\n")
        return "\n".join(lines).encode("utf-8")

    doc = Document()
    doc.add_heading(f"収益管理表 差異サマリ — {month}", 0)

    counts = {"pass": 0, "warn": 0, "fail": 0, "info": 0}
    for r in results:
        counts[r.status] = counts.get(r.status, 0) + 1

    doc.add_heading("サマリ", 1)
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    headers_row = table.rows[0].cells
    for i, h in enumerate(["総チェック数", "PASS", "要確認", "NG", "参考"]):
        headers_row[i].text = h
    values_row = table.rows[1].cells
    for i, v in enumerate([len(results), counts["pass"], counts["warn"], counts["fail"], counts["info"]]):
        values_row[i].text = str(v)

    doc.add_heading("差異一覧", 1)
    fail_warn = [r for r in results if r.status in ("fail", "warn")]
    if fail_warn:
        for r in fail_warn:
            p = doc.add_paragraph()
            run = p.add_run(f"[{r.rule_id}] {r.label}")
            run.bold = True
            if r.status == "fail":
                run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            else:
                run.font.color.rgb = RGBColor(0xFF, 0x80, 0x00)
            if r.diff is not None:
                doc.add_paragraph(f"  差異: {r.diff:+,.0f}円 ({r.diff_rate:.2%})")
            doc.add_paragraph(f"  根拠A: {r.a_source}")
            doc.add_paragraph(f"  根拠B: {r.b_source}")
            if r.note:
                doc.add_paragraph(f"  備考: {r.note}")
    else:
        doc.add_paragraph("差異・要確認項目はありません。")

    if irregulars:
        doc.add_heading("台帳イレギュラー検知", 1)
        for ir in irregulars:
            doc.add_paragraph(f"・{ir.get('description', '')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
